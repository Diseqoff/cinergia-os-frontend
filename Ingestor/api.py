import os
import shutil
import uuid
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from supabase import create_client, Client
from docx import Document
import re

# 1. Configuración Inicial
app = FastAPI(title="Cinergia OS Ingestor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 2. Conexión a Supabase (Ruta Absoluta a Prueba de Balas)
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / "credenciales.env"

load_dotenv(ENV_PATH)

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")

# Si no hay credenciales, matamos el servidor antes de que arranque a medias
if not SUPABASE_URL or not SUPABASE_KEY:
    raise RuntimeError(f"❌ ERROR FATAL: Faltan credenciales. Verifica que {ENV_PATH} exista y tenga SUPABASE_URL y SUPABASE_KEY.")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# 3. Lógica de Extracción
def procesar_documento(ruta_archivo):
    doc = Document(ruta_archivo)
    texto_completo = "\n".join([parrafo.text for parrafo in doc.paragraphs])
    
    nuevo_id = f"PRJ-{str(uuid.uuid4())[:4].upper()}"
    
    datos = {
        "id_proyecto": nuevo_id,
        "nombre": "Proyecto Sin Nombre",
        "area": "Proyectos",
        "responsable": "Sin asignar",
        "fecha": "01 Ene 2026",
        "compromisos": []
    }

    match_nombre = re.search(r"Proyecto:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_nombre: datos["nombre"] = match_nombre.group(1).strip()

    match_area = re.search(r"Área:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_area: datos["area"] = match_area.group(1).strip()

    match_resp = re.search(r"Responsable:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_resp: datos["responsable"] = match_resp.group(1).strip()

    match_fecha = re.search(r"Fecha:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_fecha: datos["fecha"] = match_fecha.group(1).strip()

    en_compromisos = False
    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        if "compromisos:" in texto.lower():
            en_compromisos = True
            continue
        
        if en_compromisos and (texto.startswith("-") or texto.startswith("•")):
            datos["compromisos"].append(texto[1:].strip())
        elif en_compromisos and texto == "":
            pass
        elif en_compromisos and not (texto.startswith("-") or texto.startswith("•")):
            en_compromisos = False

    return datos

# 4. El Endpoint
@app.post("/api/upload-acta")
async def upload_acta(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="El archivo debe ser un .docx")

    os.makedirs("documentos", exist_ok=True)
    file_path = f"documentos/{file.filename}"
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        datos = procesar_documento(file_path)
        
        supabase.table("proyectos").insert({
            "id": datos["id_proyecto"],
            "nombre": datos["nombre"],
            "area": datos["area"],
            "responsable": datos["responsable"],
            "sede": "Centro de Convenciones PE"
        }).execute()

        for comp in datos["compromisos"]:
            supabase.table("compromisos").insert({
                "proyecto_id": datos["id_proyecto"],
                "descripcion": comp
            }).execute()
            
        supabase.table("historial_etapas").insert({
            "proyecto_id": datos["id_proyecto"],
            "etapa": "Planificación Base"
        }).execute()

        os.remove(file_path)
        
        return {"status": "success", "message": "Acta procesada y subida a Cinergia OS", "data": datos}

    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)