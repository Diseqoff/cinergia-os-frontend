import os
import re
from docx import Document
from dotenv import load_dotenv
from supabase import create_client, Client

# 1. Cargar las llaves maestras
load_dotenv("Ingestor/credenciales.env")

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")

if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError("❌ Faltan credenciales. Revisa tu archivo Ingestor/credenciales.env")

# 2. Conectar a la base de datos
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
CARPETA_DOCUMENTOS = "./Ingestor/documentos/"

def extraer_datos_acta(ruta_archivo):
    doc = Document(ruta_archivo)
    texto_completo = "\n".join([parrafo.text for parrafo in doc.paragraphs])
    
    datos = {
        "id_proyecto": "PRJ-016", # ID duro para esta prueba
        "nombre": "Desconocido",
        "area": "Proyectos",
        "responsable": "Sin asignar",
        "fecha": "01 Ene 2026",
        "compromisos": []
    }

    match_nombre = re.search(r"Proyecto:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_nombre: datos["nombre"] = match_nombre.group(1).strip()

    match_area = re.search(r"Área:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_area: datos["area"] = match_area.group(1).strip()

    match_resp = re.search(r"Responsable:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_resp: datos["responsable"] = match_resp.group(1).strip()

    # CORRECCIÓN: Extracción de la fecha
    match_fecha = re.search(r"Fecha:\s*(.+)", texto_completo, re.IGNORECASE)
    if match_fecha: datos["fecha"] = match_fecha.group(1).strip()

    en_compromisos = False
    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        if "compromisos:" in texto.lower():
            en_compromisos = True
            continue
        
        if en_compromisos and (texto.startswith("-") or texto.startswith("•")):
            datos["compromisos"].append(texto[1:].strip())
        elif en_compromisos and texto == "":
            pass
        elif en_compromisos and not (texto.startswith("-") or texto.startswith("•")):
            en_compromisos = False

    return datos

def procesar_carpeta():
    print("Iniciando conexión con Supabase...\n")
    archivos = [f for f in os.listdir(CARPETA_DOCUMENTOS) if f.endswith('.docx')]
    
    for archivo in archivos:
        ruta = os.path.join(CARPETA_DOCUMENTOS, archivo)
        datos = extraer_datos_acta(ruta)
        
        print(f"📦 Subiendo datos de: {datos['nombre']}...")
        
        try:
            # A. Insertar en la tabla maestra
            supabase.table("proyectos").insert({
                "id": datos["id_proyecto"],
                "nombre": datos["nombre"],
                "area": datos["area"],
                "responsable": datos["responsable"],
                "sede": "Centro de Convenciones PE" # Dato por defecto
            }).execute()
            print("   ✅ Proyecto registrado en la BD.")

            # B. Insertar cada compromiso por separado
            for comp in datos["compromisos"]:
                supabase.table("compromisos").insert({
                    "proyecto_id": datos["id_proyecto"],
                    "descripcion": comp
                }).execute()
            print(f"   ✅ {len(datos['compromisos'])} compromisos registrados.")
            
            # C. Arrancar el reloj en el Motor Temporal
            supabase.table("historial_etapas").insert({
                "proyecto_id": datos["id_proyecto"],
                "etapa": "Planificación Base"
            }).execute()
            print("   ✅ Reloj analítico iniciado.\n")

        except Exception as e:
            print(f"   ❌ Error al insertar en BD: {e}")

if __name__ == "__main__":
    procesar_carpeta()