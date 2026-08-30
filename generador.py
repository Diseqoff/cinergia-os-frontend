from docx import Document
import os

# Asegurar que la carpeta exista
os.makedirs("Ingestor/documentos", exist_ok=True)

# Crear un documento Word real en memoria
doc = Document()

# Inyectar los datos de prueba
doc.add_paragraph("Proyecto: Seminario de Liderazgo PE")
doc.add_paragraph("Área: Eventos")
doc.add_paragraph("Responsable: Nick Huaranga")
doc.add_paragraph("Fecha: 10 Oct 2026")
doc.add_paragraph("")
doc.add_paragraph("Compromisos:")
doc.add_paragraph("- Alquilar el auditorio principal")
doc.add_paragraph("- Confirmar asistencia de ponentes")
doc.add_paragraph("- Enviar correos de marketing")

# Guardar el binario real
ruta = "Ingestor/documentos/acta_prueba.docx"
doc.save(ruta)
print(f"Archivo Word real generado con éxito en: {ruta}")